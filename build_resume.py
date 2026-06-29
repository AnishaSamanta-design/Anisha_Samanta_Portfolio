#!/usr/bin/env python3
"""Generate an ATS-friendly UI/UX Designer resume as .docx and .pdf."""

import os

# ── Content ───────────────────────────────────────────────────────────────
NAME = "Anisha Samanta"
TITLE = "UI/UX & Product Designer"
PHONE = "+91-7908912952"
EMAIL = "anishasamanta331@gmail.com"
LINKS = [
    ("LinkedIn", "https://www.linkedin.com/in/anisha-samanta-professional/"),
    ("Behance", "https://www.behance.net/anishasamanta27"),
    ("Portfolio", "https://anisha-samanta-portfolio.vercel.app/"),
]

SUMMARY = (
    "Product Designer with 3 years of experience designing enterprise applications across finance, healthcare, "
    "retail, transportation, e-commerce and others domains. Backed by 5 years in operations engineering, I bring "
    "a strong understanding of complex workflows and user challenges. I focus on creating clear, practical "
    "experiences that help users complete tasks efficiently while supporting business goals."
)

SKILLS = [
    ("Design Tools", "Figma, Miro, Framer"),
    ("UX & Product Design", "UX Research, Product Thinking, User Flows, Information Architecture, "
                            "Interaction Design, Wireframing, Prototyping, Journey Mapping, Accessibility, Design Systems"),
    ("Technical Knowledge", "C, C++, HTML Fundamentals"),
    ("Languages", "English, Bengali, Hindi"),
]

RELEVANT_TITLE = "UI/UX Designer"
RELEVANT_ORG = "Accenture"
RELEVANT_DATES = "Apr 2023 – Jun 2026"
RELEVANT_BULLETS = [
    "Designed and shipped responsive web experiences across different industries for both web and mobile, balancing user needs and business goals.",
    "Designed agentic, Gen AI-powered solution prototypes for 20+ clients, translating complex AI workflows into clear, intuitive interfaces.",
    "Simplified complex finance workflows by redesigning enterprise applications for RTR, OTC, FP&A, Cash Collections, and Reconciliation processes, improving usability and task efficiency.",
    "Leveraged Figma Make (AI) to accelerate delivery, contributing to new client acquisitions and multi-million-dollar contract renewals.",
    "Improved analytics and reporting interfaces, reducing dependency on support teams and improving data clarity for end users.",
    "Maintained a unified design system and component library across Agile sprints, ensuring brand consistency and cutting production time by ~40%.",
    "Created a reusable Google Design Sprint-inspired ideation template, standardizing brainstorming and concept validation across projects.",
]

OTHER_EXPERIENCE = [
    ("Security Analyst", "Accenture", "Feb 2020 – Mar 2023",
     "Ensured secure and seamless access management by administering user identities, provisioning accounts, "
     "and supporting enterprise IAM operations across global systems."),
    ("Software Asset Management", "Nexwave", "Apr 2018 – Jan 2020",
     "Optimized software asset governance by managing licenses, compliance, and lifecycle operations, ensuring "
     "cost-effective and seamless access to enterprise tools."),
]

EDUCATION = [
    ("B. Tech – Electronics & Communication Engineering", "Supreme Knowledge Foundation Group of Institutions, 2013–2017"),
    ("Advanced Certification in UI/UX Design", "IIIT-B, 2024"),
    ("Design Thinking Define & Ideate", "Accenture, Feb 2026"),
]

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUT_DIR, "Anisha_Samanta_Resume2026.docx")
PDF_PATH = os.path.join(OUT_DIR, "Anisha_Samanta_Resume2026.pdf")
PDF_COPIES = [os.path.join(OUT_DIR, "Anisha_Samanta_Portfolio", "Anisha_Samanta_Resume2026.pdf")]

# Sophisticated earthy palette — warm, editorial, brand-aligned
ACCENT = "A9542F"   # terracotta — accent rule, role, hyperlinks
INK = "2B2622"      # warm charcoal — name, headings, roles, companies
MUTED = "8A6F52"    # warm taupe — title & supporting text
SAND = "EFE7D9"     # soft sand — header masthead background
RULE = "DCD2C2"     # warm hairline
NAVY = ACCENT       # back-compat alias (older refs)

# ── DOCX ──────────────────────────────────────────────────────────────────
def build_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    def add_hyperlink(paragraph, url, text, color="1F3864", size_hp="18"):
        r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), size_hp); rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    doc = Document()
    # ATS-friendly base style: standard font, single column, normal margins
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    def set_space(p, before=0, after=0, line=1.08):
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line

    def section_heading(text):
        p = doc.add_paragraph()
        set_space(p, before=16, after=5)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(INK)
        # letter-spacing for a refined tracked-caps look
        rPr = run._element.get_or_add_rPr()
        spc = OxmlElement("w:spacing"); spc.set(qn("w:val"), "30"); rPr.append(spc)
        # thin teal underline accent
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), ACCENT)
        pbdr.append(bottom)
        pPr.append(pbdr)

    # Header — editorial left-aligned masthead on a soft sand panel
    def header_panel(p, space_after):
        set_space(p, after=space_after)
        p.paragraph_format.left_indent = Inches(0.13)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), SAND)
        pPr.append(shd)
        pbdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "22")
        left.set(qn("w:space"), "10"); left.set(qn("w:color"), ACCENT)
        pbdr.append(left)
        pPr.append(pbdr)

    p = doc.add_paragraph(); header_panel(p, 2)
    r = p.add_run(NAME.upper())
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(INK)
    rPr = r._element.get_or_add_rPr()
    spc = OxmlElement("w:spacing"); spc.set(qn("w:val"), "50"); rPr.append(spc)

    p = doc.add_paragraph(); header_panel(p, 5)
    r = p.add_run(TITLE.upper())
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    rPr = r._element.get_or_add_rPr()
    spc = OxmlElement("w:spacing"); spc.set(qn("w:val"), "40"); rPr.append(spc)

    # Contact + links on one line
    p = doc.add_paragraph(); header_panel(p, 2)
    r = p.add_run(f"{PHONE}   ·   {EMAIL}      ")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("6E6258")
    for i, (label, url) in enumerate(LINKS):
        if i:
            sep = p.add_run("   ·   ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = RGBColor.from_string("6E6258")
        add_hyperlink(p, url, label, color=ACCENT)

    # Summary
    section_heading("Professional Summary")
    p = doc.add_paragraph()
    set_space(p, after=2, line=1.18)
    p.add_run(SUMMARY)

    # Skills
    section_heading("Core Skills")
    for label, val in SKILLS:
        p = doc.add_paragraph()
        set_space(p, after=2, line=1.15)
        rb = p.add_run(label + ": ")
        rb.bold = True
        p.add_run(val)

    # Work experience
    section_heading("Work Experience")

    p = doc.add_paragraph()
    set_space(p, before=4, after=3)
    r = p.add_run("—  Relevant Experience")
    r.bold = True
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p = doc.add_paragraph()
    set_space(p, after=1)
    r = p.add_run(RELEVANT_TITLE)
    r.bold = True
    r.font.size = Pt(11)
    p.add_run("   ·   ").font.color.rgb = RGBColor.from_string("AAAAAA")
    co = p.add_run(RELEVANT_ORG)
    co.italic = True
    co.font.color.rgb = RGBColor.from_string(INK)
    d = p.add_run(f"      {RELEVANT_DATES}")
    d.italic = True
    d.font.size = Pt(9.5)
    d.font.color.rgb = RGBColor.from_string("777777")

    for b in RELEVANT_BULLETS:
        p = doc.add_paragraph(style="List Bullet")
        set_space(p, after=4, line=1.16)
        p.add_run(b)

    p = doc.add_paragraph()
    set_space(p, before=7, after=3)
    r = p.add_run("—  Other Work Experience")
    r.bold = True
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    for title, org, dates, desc in OTHER_EXPERIENCE:
        p = doc.add_paragraph()
        set_space(p, before=2, after=1)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run("   ·   ").font.color.rgb = RGBColor.from_string("AAAAAA")
        co = p.add_run(org)
        co.italic = True
        co.font.color.rgb = RGBColor.from_string(INK)
        d = p.add_run(f"      {dates}")
        d.italic = True
        d.font.size = Pt(9.5)
        d.font.color.rgb = RGBColor.from_string("777777")
        p = doc.add_paragraph(style="List Bullet")
        set_space(p, after=4, line=1.16)
        p.add_run(desc)

    # Education
    section_heading("Education & Certifications")
    for title, detail in EDUCATION:
        p = doc.add_paragraph(style="List Bullet")
        set_space(p, after=4, line=1.16)
        r = p.add_run(title)
        r.bold = True
        p.add_run(" — " + detail)

    doc.save(DOCX_PATH)
    print("Saved:", DOCX_PATH)


# ── PDF ───────────────────────────────────────────────────────────────────
def build_pdf():
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, ListFlowable,
                                     ListItem, HRFlowable)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    accent = colors.HexColor("#" + ACCENT)
    ink = colors.HexColor("#" + INK)
    rule = colors.HexColor("#" + RULE)
    muted = colors.HexColor("#" + MUTED)
    sand = colors.HexColor("#" + SAND)
    gray = colors.HexColor("#6E6258")   # warm gray
    styles = getSampleStyleSheet()

    HI = 15  # header indent — clears the vertical accent rule
    HGAP = 10  # uniform gap (pt ≈ 13px) between header lines
    name_style = ParagraphStyle("name", parent=styles["Normal"], fontName="Helvetica-Bold",
                                fontSize=26, textColor=ink, leading=27, spaceAfter=HGAP,
                                alignment=TA_LEFT, leftIndent=HI)
    title_style = ParagraphStyle("title", parent=styles["Normal"], fontName="Helvetica-Bold",
                                  fontSize=9.5, textColor=accent, leading=11, spaceAfter=HGAP,
                                  alignment=TA_LEFT, leftIndent=HI)
    contact_style = ParagraphStyle("contact", parent=styles["Normal"], fontName="Helvetica",
                                    fontSize=8.5, leading=11, spaceAfter=HGAP,
                                    alignment=TA_LEFT, leftIndent=HI, textColor=gray)
    sect_style = ParagraphStyle("sect", parent=styles["Normal"], fontName="Helvetica-Bold",
                                fontSize=10.5, textColor=ink, leading=12, spaceBefore=4, spaceAfter=2)
    sub_style = ParagraphStyle("sub", parent=styles["Normal"], fontName="Helvetica-BoldOblique",
                               fontSize=9, textColor=accent, leading=11, spaceBefore=5, spaceAfter=3)
    body_style = ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica",
                                fontSize=9.5, leading=13.5, alignment=TA_LEFT, spaceAfter=2)
    role_style = ParagraphStyle("role", parent=styles["Normal"], fontName="Helvetica",
                                fontSize=10, leading=13, spaceAfter=2)
    bullet_style = ParagraphStyle("bullet", parent=styles["Normal"], fontName="Helvetica",
                                  fontSize=9.5, leading=13.5, spaceAfter=1.5)

    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def section(text):
        # uppercase heading + a short rust accent bar (modern, ATS-safe)
        return [Paragraph(text.upper(), sect_style),
                HRFlowable(width=44, thickness=2.4, color=accent, hAlign="LEFT",
                           spaceBefore=1, spaceAfter=4)]

    def bullets(items):
        return ListFlowable(
            [ListItem(Paragraph(t, bullet_style), leftIndent=12, value="•") for t in items],
            bulletType="bullet", start="•", leftIndent=10, bulletFontSize=8,
            bulletColor=accent, spaceBefore=0, spaceAfter=0,
        )

    story = []
    story.append(Paragraph(NAME.upper(), name_style))
    story.append(Paragraph(esc(TITLE).upper(), title_style))
    story.append(Paragraph(
        f"{esc(PHONE)}&nbsp;&nbsp;&#183;&nbsp;&nbsp;{esc(EMAIL)}", contact_style))
    link_html = "&nbsp;&nbsp;&#183;&nbsp;&nbsp;".join(
        f'<a href="{u}" color="#{ACCENT}"><u>{esc(t)}</u></a>' for t, u in LINKS)
    story.append(Paragraph(link_html, contact_style))
    story.append(Spacer(1, 4))  # links already carries HGAP below it; small extra before body

    story += section("Professional Summary")
    story.append(Paragraph(esc(SUMMARY), body_style))

    story += section("Core Skills")
    for label, val in SKILLS:
        story.append(Paragraph(f"<b>{esc(label)}:</b> {esc(val)}", body_style))

    story += section("Work Experience")
    story.append(Paragraph("—&nbsp;&nbsp;Relevant Experience", sub_style))
    story.append(Paragraph(
        f"<b>{esc(RELEVANT_TITLE)}</b>"
        f"<font color='#AAAAAA'>&nbsp;&nbsp;&#183;&nbsp;&nbsp;</font>"
        f"<i>{esc(RELEVANT_ORG)}</i>"
        f"&nbsp;&nbsp;&nbsp;&nbsp;<font color='#777777' size=9><i>{RELEVANT_DATES}</i></font>",
        role_style))
    story.append(bullets([esc(b) for b in RELEVANT_BULLETS]))

    story.append(Paragraph("—&nbsp;&nbsp;Other Work Experience", sub_style))
    for title, org, dates, desc in OTHER_EXPERIENCE:
        story.append(Paragraph(
            f"<b>{esc(title)}</b>"
            f"<font color='#AAAAAA'>&nbsp;&nbsp;&#183;&nbsp;&nbsp;</font>"
            f"<i>{esc(org)}</i>"
            f"&nbsp;&nbsp;&nbsp;&nbsp;<font color='#777777' size=9><i>{dates}</i></font>",
            role_style))
        story.append(bullets([esc(desc)]))

    story += section("Education & Certifications")
    edu_items = [f"<b>{esc(t)}</b> — {esc(d)}" for t, d in EDUCATION]
    story.append(bullets(edu_items))

    # Editorial masthead drawn behind the real header text (ATS-safe — text stays selectable).
    # The header flowables (name → links) are vertically CENTERED inside the sand band, and the
    # terracotta vertical rule spans that same block — so all details sit on one centred axis.
    # visible block = name(27) + gap + title(11) + gap + contact(11) + gap + links(11)
    BLOCK = 27 + 11 + 11 + 11 + 3 * 10   # = 90 pt
    band_h = 1.42 * inch
    LM = 0.7 * inch
    TOPM = (band_h - BLOCK) / 2.0   # equal sand margin above the name and below the links

    def draw_header_band(canvas, doc_):
        w, h = doc_.pagesize
        canvas.saveState()
        # soft sand panel, full bleed
        canvas.setFillColor(sand)
        canvas.rect(0, h - band_h, w, band_h, fill=1, stroke=0)
        # warm hairline at the panel base
        canvas.setFillColor(rule)
        canvas.rect(0, h - band_h, w, 0.8, fill=1, stroke=0)
        # vertical terracotta rule spanning the centred header block
        canvas.setFillColor(accent)
        canvas.rect(LM, h - TOPM - BLOCK + 3, 3.4, BLOCK - 6, fill=1, stroke=0)
        canvas.restoreState()

    doc = SimpleDocTemplate(PDF_PATH, pagesize=letter,
                            topMargin=TOPM, bottomMargin=0.38 * inch,
                            leftMargin=LM, rightMargin=0.7 * inch,
                            title="Anisha Samanta — UI/UX Designer Resume", author=NAME)
    doc.build(story, onFirstPage=draw_header_band)
    print("Saved:", PDF_PATH)


if __name__ == "__main__":
    import shutil
    build_docx()
    build_pdf()
    for dest in PDF_COPIES:
        if os.path.isdir(os.path.dirname(dest)):
            shutil.copyfile(PDF_PATH, dest)
            print("Copied to:", dest)
